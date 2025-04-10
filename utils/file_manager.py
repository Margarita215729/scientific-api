"""
utils/file_manager.py

Модуль для работы с файловой системой. Предоставляет функции для:
- Рекурсивного просмотра файлов в директории
- Чтения содержимого файлов различных типов (txt, CSV, JSON, ipynb, docx)
- Разбиения больших файлов на части
- Создания новых файлов и редактирования существующих

Дополнительно используется библиотека `python-docx` для работы с DOCX.
"""

import os
import json
import pandas as pd

# Для работы с docx
try:
    from docx import Document
except ImportError:
    Document = None

def list_files(directory: str) -> dict:
    """
    Рекурсивно просматривает директорию и возвращает вложенную структуру файлов и папок.
    
    :param directory: Корневая директория для просмотра.
    :return: Словарь, представляющий дерево файлов.
    """
    tree = {}
    for root, dirs, files in os.walk(directory):
        # Вычисляем относительный путь от корневой директории
        rel_path = os.path.relpath(root, directory)
        tree[rel_path] = {"dirs": dirs, "files": files}
    return tree

def read_file(file_path: str) -> str:
    """
    Читает содержимое файла и возвращает его в виде строки.
    Поддерживаются: txt, ipynb, JSON. Для CSV возвращает строковое представление.
    Если файл DOCX, и библиотека docx установлена, читает все параграфы.
    
    :param file_path: Путь к файлу.
    :return: Строковое содержимое файла.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext in [".txt", ".ipynb"]:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif ext == ".json":
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return json.dumps(data, indent=2)
    elif ext == ".csv":
        df = pd.read_csv(file_path)
        return df.to_csv(index=False)
    elif ext == ".docx":
        if Document is None:
            raise ImportError("Для работы с DOCX установите пакет python-docx")
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        with open(file_path, "rb") as f:
            return f.read().hex()  # Если бинарный тип, возвращаем hex-представление

def split_file(file_path: str, chunk_size: int = 1000) -> list:
    """
    Разбивает текстовый файл на части заданного размера (в символах).
    
    :param file_path: Путь к файлу.
    :param chunk_size: Размер каждой части (по умолчанию 1000 символов).
    :return: Список строк – частей файла.
    """
    content = read_file(file_path)
    # Если содержимое не строка, пытаемся декодировать
    if not isinstance(content, str):
        content = content.decode("utf-8", errors="ignore")
    return [content[i:i + chunk_size] for i in range(0, len(content), chunk_size)]

def create_file(file_path: str, content: str, file_type: str = "txt") -> None:
    """
    Создаёт новый файл с указанным содержимым.
    Поддерживает: txt, JSON, CSV, ipynb, docx.
    
    :param file_path: Путь, по которому будет создан файл.
    :param content: Содержимое файла в виде строки (для CSV/JSON можно передавать json-строку).
    :param file_type: Тип файла (txt, json, csv, ipynb, docx).
    """
    ext = file_type.lower()
    if ext == "txt":
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
    elif ext == "json":
        data = json.loads(content)
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2)
    elif ext == "csv":
        # Преобразование через Pandas: ожидается, что content – CSV строка
        from io import StringIO
        df = pd.read_csv(StringIO(content))
        df.to_csv(file_path, index=False)
    elif ext == "ipynb":
        # ipynb – JSON формат
        data = json.loads(content)
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2)
    elif ext == "docx":
        if Document is None:
            raise ImportError("Для создания DOCX установите пакет python-docx")
        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc.save(file_path)
    else:
        raise ValueError(f"Неподдерживаемый тип файла: {file_type}")

def edit_file(file_path: str, new_content: str) -> None:
    """
    Редактирует существующий файл, перезаписывая его новым содержимым.
    Функция определяет тип файла по расширению и сохраняет содержимое в текстовом формате.
    
    :param file_path: Путь к существующему файлу.
    :param new_content: Новое содержимое файла.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext in [".txt", ".ipynb", ".json", ".csv"]:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(new_content)
    elif ext == ".docx":
        if Document is None:
            raise ImportError("Для работы с DOCX установите пакет python-docx")
        doc = Document()
        for line in new_content.split("\n"):
            doc.add_paragraph(line)
        doc.save(file_path)
    else:
        # Если файл бинарный или неизвестного формата, сохраняем как текст (можно расширить логику по необходимости)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(new_content)
